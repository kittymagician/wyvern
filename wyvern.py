import requests
import docx
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pendulum
from datetime import datetime
import shodan
import openai

wyvern_art = '''
                /\\
               //\\\\
      |\\___/|//  \\\\        /\\
      /0  0  `      \\      //\\\\
     (   /   =\\  /=\\ \\   ///\\\\\\\\   
      \\\\/     //\\\\//\\\\\\/  |   \\\\\\\\\\  
       \\\\    //  \\\\\\\\\\\\ \\_ _\\\\_\\\\_/
       `\\"`   `\\"`  `\\"``   `\\"`\\"` 

       Github: github.com/kittymagician/wyvern
       Author: kittymagician
       Licence: MIT License
'''

print(wyvern_art)

# Define timezone
utc = pendulum.timezone('UTC')

# Define the domain to query
domain = "domain goes here"
# Define company name for Shodan
company_name = "company name goes here"
# Define if Shodan query should run
noShodan = False
# Define the DNS record types to query
record_types = ["A", "MX", "TXT"]

# Query the DNS records using Google DNS over HTTPS
results = {}
spf_records = []
dmarc_records = []
for record_type in record_types:
    url = f"https://dns.google/resolve?name={domain}&type={record_type}"
    response = requests.get(url).json()
    if "Answer" in response:
        if record_type == "TXT":
            for answer in response["Answer"]:
                if answer["data"].startswith("v=spf1"):
                    spf_records.append(answer["data"])
                if answer["data"].startswith("v=DMARC1"):
                    dmarc_records.append(answer["data"])
                else:
                  if "TXT" not in results:
                    results["TXT"] = []
                  results["TXT"].append(answer["data"])

        else:
          if record_type not in results:
            results[record_type] = []
          results[record_type].extend([answer["data"] for answer in response["Answer"]])

# Define your OpenAI API token
openai.api_key = ""

# Define your ipinfo.io API token
ipinfo_api_token = ""

# Define your Shodan API key
shodan_api_key = ""
api = shodan.Shodan(shodan_api_key)

# Get ASN details for A Records
if "A" in results:
    asn_details = []
    for ip in results["A"]:
        response = requests.get(f"https://ipinfo.io/{ip}?token={ipinfo_api_token}").json()
        asn = response.get("org", "N/A")
        asn_details.append(asn)

# Get IP information for MX Records
if "MX" in results:
    mx_info = []
    for mx in results["MX"]:
        mx_domain = mx.split()[-1]
        response = requests.get(f"https://dns.google/resolve?name={mx_domain}&type=A").json()
        if "Answer" in response:
            ip = response["Answer"][0]["data"]
            response = requests.get(f"https://ipinfo.io/{ip}?token={ipinfo_api_token}").json()
            ip_info = response.get("org", "N/A")
            mx_info.append((mx, ip, ip_info))

# Try to obtain Shodan Data
vulnerable_hosts = []
try:
    query = f'org:"{company_name}"'
    shodan_results = api.search(query)
    if shodan_results['total'] == 0:
        print(f"No hosts found for {company_name}.")
        noShodan = True
    else:
        for result in shodan_results['matches']:
            host_ip = result['ip_str'] if 'ip_str' in result else 'N/A'
            hostnames = result['hostnames'][0] if 'hostnames' in result and result['hostnames'] else 'N/A'
            cve_list = result['vulns'] if 'vulns' in result else {}
            cve_text = "\n".join([f"{cve}: {cve_info['summary']}" for cve, cve_info in cve_list.items()])
            if cve_text:
                vulnerable_hosts.append({'ip': host_ip, 'hostname': hostnames, 'cves': cve_text})
except shodan.APIError as e:
    print(f"Error: {e}")

# Function to add table borders
def add_table_borders(cell):
    tc = cell._element.tcPr
    tcBorders = OxmlElement('w:tcBorders')
    for border_type in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f"w:{border_type}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tc.append(tcBorders)

# Create a new Word document
doc = docx.Document()

# Add a title page
doc.add_heading(f"DNS Records for {domain}", 0)
dt = datetime.now(utc)
doc.add_paragraph(f"{dt}")

# Add a page break
doc.add_page_break()

doc.add_paragraph(f"The following DNS records were retrieved for {domain}:")
for record_type in record_types:
    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.add_run(f"{record_type} Records").bold = True

doc.add_page_break()

# Add record types to ToC
doc.add_heading("Table of Contents", level=1)
toc = doc.add_paragraph()
for record_type in ["A", "MX", "CNAME", "TXT"]:
    if record_type in results:
        toc.add_run(f"{record_type} Records\n")
        toc.add_run("\n")

# Add the DNS records to the document
doc.add_page_break()

for record_type, values in results.items():
    doc.add_heading(f"{record_type} Records", level=1)
    if len(values) > 0:
        if record_type == "A":
            table = doc.add_table(rows=len(values), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, value in enumerate(values):
                cell = table.cell(i, 0)
                cell.text = value
                cell.width = Cm(15)
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_table_borders(cell)
                cell = table.cell(i, 1)
                cell.text = asn_details[i]
                cell.width = Cm(15)
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_table_borders(cell)
            doc.add_paragraph(f"A Records are IP Addresses that point to a domain name.")
        elif record_type == "MX" and mx_info:
            table = doc.add_table(rows=len(mx_info), cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (mx, ip, ip_info) in enumerate(mx_info):
                cell = table.cell(i, 0)
                cell.text = mx
                cell.width = Cm(15)
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_table_borders(cell)
                cell = table.cell(i, 1)
                cell.text = ip
                cell.width = Cm(15)
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_table_borders(cell)
                cell = table.cell(i, 2)
                cell.text = ip_info
                cell.width = Cm(15)
                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_table_borders(cell)
            doc.add_paragraph(f"MX Records specify the mail servers responsible for accepting email messages on behalf of a domain name.")
        else:
            table = doc.add_table(rows=len(values), cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, value in enumerate(values):
              cell = table.cell(i, 0)
              cell.text = value
              cell.width = Cm(15)
              cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
              add_table_borders(cell)

# Add SPF record data if available
if spf_records:
    doc.add_heading("SPF Record", level=1)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = spf_records[0]
    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.width = Cm(16)
    cell.height = Cm(1.5)

# Add DMARC record data if available
if dmarc_records:
    doc.add_heading("DMARC Record", level=1)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = dmarc_records[0]
    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.width = Cm(16)
    cell.height = Cm(1.5)
doc.add_page_break()
# Add Shodan data if there are vulnerable hosts
if vulnerable_hosts:
    doc.add_heading("Shodan Data", level=1)
    table = doc.add_table(rows=len(vulnerable_hosts) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Add header row
    header_row = table.rows[0]
    header_row.cells[0].text = "IP Address"
    header_row.cells[1].text = "Hostname"
    header_row.cells[2].text = "Vulnerabilities"
    add_table_borders(header_row.cells[0])
    add_table_borders(header_row.cells[1])
    add_table_borders(header_row.cells[2])

    # Add data rows
    for i, host in enumerate(vulnerable_hosts):
        row = table.rows[i+1]
        row.cells[0].text = host['ip']
        row.cells[1].text = host['hostname'] if 'hostname' in host else 'N/A'
        row.cells[2].text = host['cves']
        add_table_borders(row.cells[0])
        add_table_borders(row.cells[1])
        add_table_borders(row.cells[2])
    doc.add_page_break()

# OpenAI Summery
def generate_summary(prompt, token):
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=token,
        n=1,
        stop=None,
        temperature=0.7,
    )

    message = response.choices[0].text.strip()
    return message

report_text = f"DNS Records for {domain}\n\n"
for record_type, values in results.items():
    report_text += f"{record_type} Records:\n"
    if len(values) > 0:
        for value in values:
            report_text += f"{value}\n"
    report_text += "\n"
if vulnerable_hosts:
    report_text += "Shodan Data:\n"
    for host in vulnerable_hosts:
        report_text += f"IP: {host['ip']}\nHostname: {host['hostname']}\nVulnerabilities:\n{host['cves']}\n\n"
print(report_text)
if noShodan is False:
  prompt = f"Please provide an overview of the findings for the following DNS records and Shodan data in 200 words or less include CVEs and provide recommendations on how to remediate:\n{report_text}"
  summary = generate_summary(prompt, 300)
  doc.add_heading("Findings Overview (OpenAi Experimental)", level=1)
  doc.add_paragraph(summary)
  doc.add_page_break()
  doc.add_heading("Remmediation Plan (OpenAi Experimental)", level=1)
  prompt = f"write a remmediation plan in 500 words or less.\n{report_text}"
  remmediation = generate_summary(prompt, 500)
  if spf_records:
    doc.add_paragraph(remmediation)
    prompt2 = f"look at the txt records for SPF, Configured securely? if not explain best practice:{spf_records[0]}"
    remmediation_2 = generate_summary(prompt2, 1500)
    doc.add_paragraph(remmediation_2)
  doc.add_page_break()
# Save the document
doc.save(f"wyvern_{domain}_{dt}.docx")